#!/usr/bin/python3
# -*- coding: utf-8 -*-
# -*- mode: python; python-indent-offset: 4 -*-
#
# ./get_textbox_submissions_as_docx.py course_id assignment_id
# Purpose:
#   create a DOCX file for each textual submission
#
# with the option "-v" or "--verbose" you get lots of output - showing in detail the operations of the program
# with the option '-C'or '--containers' use HTTP rather than HTTPS for access to Canvas
#
# Can also be called with an alternative configuration file:
#  --config config-test.json
#
# Example - to generate docx file locally
# ./get_textbox_submissions_as_docx.py -C 7 44^
#
# # Example - to generate docx file locally and upload it as a submission
# ./get_textbox_submissions_as_docx.py --submit  -C 7 44
#
# to get the DOCX library do:
#    pip3 install python-docx
# 
# G. Q. Maguire Jr.
#
# 2020.03.01
#

import requests, time
import pprint
import optparse
import sys
import json

# use the Docment library to be able to create the .docx file
from docx import Document
#from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_BREAK
from docx.shared import Pt

import os

from bs4 import BeautifulSoup

#############################
###### EDIT THIS STUFF ######
#############################

global baseUrl	# the base URL used for access to Canvas
global header	# the header for all HTML requests
global payload	# place to store additionally payload when needed for options to HTML requests

# Based upon the options to the program, initialize the variables used to access Canvas gia HTML requests
def initialize(options):
    global baseUrl, header, payload

    # styled based upon https://martin-thoma.com/configuration-files-in-python/
    if options.config_filename:
        config_file=options.config_filename
    else:
        config_file='config.json'

    try:
        with open(config_file) as json_data_file:
            configuration = json.load(json_data_file)
            access_token=configuration["canvas"]["access_token"]
            if options.containers:
                baseUrl="http://"+configuration["canvas"]["host"]+"/api/v1"
                print("using HTTP for the container environment")
            else:
                baseUrl="https://"+configuration["canvas"]["host"]+"/api/v1"

            header = {'Authorization' : 'Bearer ' + access_token}
            payload = {}
    except:
        print("Unable to open configuration file named {}".format(config_file))
        print("Please create a suitable configuration file, the default name is config.json")
        sys.exit()

def enable_docx_submission_for_assignment(course_id, assignment_id, assignment):
    global Verbose_Flag
    # Use the Canvas API to submit to edit an assignment
    #PUT /api/v1/courses/:course_id/assignments/:id
    url = "{0}/courses/{1}/assignments/{2}".format(baseUrl, course_id, assignment_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    something_to_edit=False
    # add 'online_upload' as a submission type if it is not there
    current_submission_types=assignment.get('submission_types', [])
    if not current_submission_types:
        print("Assignment does not permit submissions!")
        return False

    if 'online_upload' not in current_submission_types:
        current_submission_types.append('online_upload')
        parameters['assignment[submission_types][]']=current_submission_types
        something_to_edit=True

    # if there are specified allowed extensions, make sure they include docx, else add it to the list of extensions
    currently_allowed_extensions=assignment.get('allowed_extensions', False)
    if currently_allowed_extensions:
        if 'docx' not in currently_allowed_extensions:
            if len(currently_allowed_extensions) > 0:
                currently_allowed_extensions.append('docx')
                parameters['assignment[allowed_extensions][]']=currently_allowed_extensions
                something_to_edit=True

    if something_to_edit:
        r = requests.put(url, params=parameters, headers = header)
        if Verbose_Flag:
            print("r.status_code={}".format(r.status_code))
        if r.status_code < int(300):
            page_response=r.json()
            return page_response
    return True

def submit_file_for_assignment(course_id, assignment_id, user_id, filename):
    global Verbose_Flag
    if Verbose_Flag:
        print("in submit_file_for_assignment course_id={0}, assignment_id={1}, user_id={2}, filanem={3}".format(course_id, assignment_id, user_id, filename))
    # Use the Canvas API to submit a file for an assignment
    #POST /api/v1/courses/:course_id/assignments/:assignment_id/submissions/:user_id/files

    url = "{0}/courses/{1}/assignments/{2}/submissions/{3}/files".format(baseUrl, course_id, assignment_id, user_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    extra_parameters={'as_user_id': user_id}

    payload={'name': filename,
             'size': os.path.getsize(filename),
             'content-type': 'binary/octet-stream',
             'on_duplicate': 'overwrite'
    }
    r = requests.post(url, params=extra_parameters, headers = header, data=payload)
    if Verbose_Flag:
        print("result of post: {}".format(r.text))
        print("r.status_code={}".format(r.status_code))
    #if r.status_code == requests.codes.created:
    if r.status_code < int(300):
        page_response=r.json()
        print("first step of uploading a file for a submission")
        # {
        #   "upload_url": "https://some-bucket.s3.amazonaws.com/",
        #   "upload_params": {
        #     "key": "/users/1234/files/profile_pic.jpg",
        #     <unspecified parameters; key above will not necesarily be present either>
        #   }
        # }
        upload_url=page_response['upload_url']
        upload_params=page_response['upload_params']
        key=page_response['upload_params']['key']
        print("upload_url={0}, key={1}, upload_params={2}".format(upload_url,key, upload_params))

        r = requests.post(upload_url,params=upload_params, files={"file": open(filename, 'rb')})
        if r.status_code < int(300):
            page_response=r.json()
            print("page_response={}".format(page_response))
            file_id=page_response['id']
            print("file_id={}".format(file_id))

            # POST /api/v1/courses/:course_id/assignments/:assignment_id/submissions
            url = "{0}/courses/{1}/assignments/{2}/submissions".format(baseUrl, course_id, assignment_id)
            if Verbose_Flag:
                print("url: {}".format(url))

            extra_parameters={'as_user_id': user_id,
                              'comment[text_comment]': "uploaded DOCX file",
                              'submission[submission_type]': 'online_upload',
                              'submission[file_ids][]': file_id,
                              'submission[user_id]': user_id
                              }

            r = requests.post(url, params=extra_parameters, headers = header)
            print("r.status_code={}".format(r.status_code))
            r.raise_for_status()
            if r.status_code < int(300):
                page_response=r.json()
                print("page_response={}".format(page_response))
                return page_response
    return False

def get_paragraphs_from_HTML(b):
    # clean up the HTML
    b.replace('<br>', '</p><p>')

    list_of_paragraphs=[]
    xml=BeautifulSoup(b, "lxml")
    paragraphs=xml.findAll('p')
    for p in paragraphs:
        list_of_paragraphs.append(p)

    return list_of_paragraphs

    
def get_user(user_id):
    #GET /api/v1/users/:user_id
    url = "{0}/users/{1}".format(baseUrl,user_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    r = requests.get(url, headers = header)
    if r.status_code == requests.codes.ok:
        page_response=r.json()

        return page_response

    return []

def users_in_course(course_id):
    user_found_thus_far=[]
    # Use the Canvas API to get the list of users enrolled in this course
    #GET /api/v1/courses/:course_id/enrollments

    url = "{0}/courses/{1}/enrollments".format(baseUrl,course_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    extra_parameters={'per_page': '100'}
    r = requests.get(url, params=extra_parameters, headers = header)
    if Verbose_Flag:
        print("result of getting enrollments: {}".format(r.text))

    if r.status_code == requests.codes.ok:
        page_response=r.json()

        for p_response in page_response:  
            user_found_thus_far.append(p_response)

        # the following is needed when the reponse has been paginated
        # i.e., when the response is split into pieces - each returning only some of the list of modules
        # see "Handling Pagination" - Discussion created by tyler.clair@usu.edu on Apr 27, 2015, https://community.canvaslms.com/thread/1500
        while r.links.get('next', False):
            r = requests.get(r.links['next']['url'], headers=header)  
            page_response = r.json()  
            if r.status_code == requests.codes.ok:
                for p_response in page_response:  
                    user_found_thus_far.append(p_response)

    return user_found_thus_far

def list_assignments(course_id):
    assignments_found_thus_far=[]
    # Use the Canvas API to get the list of assignments for the course
    #GET /api/v1/courses/:course_id/assignments

    url = "{0}/courses/{1}/assignments".format(baseUrl, course_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    r = requests.get(url, headers = header)
    if Verbose_Flag:
        print("result of getting assignments: {}".format(r.text))

    if r.status_code == requests.codes.ok:
        page_response=r.json()

        for p_response in page_response:  
            assignments_found_thus_far.append(p_response)

        # the following is needed when the reponse has been paginated
        # i.e., when the response is split into pieces - each returning only some of the list of assignments
        # see "Handling Pagination" - Discussion created by tyler.clair@usu.edu on Apr 27, 2015, https://community.canvaslms.com/thread/1500
        while r.links['current']['url'] != r.links['last']['url']:  
            r = requests.get(r.links['next']['url'], headers=header)  
            if Verbose_Flag:
                print("result of getting assignments for a paginated response: {}".format(r.text))
            page_response = r.json()  
            for p_response in page_response:  
                assignments_found_thus_far.append(p_response)

    return assignments_found_thus_far

def list_ungraded_assignments(course_id):
    assignments_found_thus_far=[]
    # Use the Canvas API to get the list of assignments for the course
    #GET /api/v1/courses/:course_id/assignments

    url = "{0}/courses/{1}/assignments".format(baseUrl, course_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    extra_parameters={'bucket': 'ungraded'
                      }
    r = requests.get(url, params=extra_parameters, headers = header)

    if Verbose_Flag:
        print("result of getting ungraded assignments: {}".format(r.text))

    if r.status_code == requests.codes.ok:
        page_response=r.json()

        for p_response in page_response:  
            assignments_found_thus_far.append(p_response)

        # the following is needed when the reponse has been paginated
        # i.e., when the response is split into pieces - each returning only some of the list of assignments
        # see "Handling Pagination" - Discussion created by tyler.clair@usu.edu on Apr 27, 2015, https://community.canvaslms.com/thread/1500
        while r.links['current']['url'] != r.links['last']['url']:  
            r = requests.get(r.links['next']['url'], headers=header)  
            if Verbose_Flag:
                print("result of getting assignments for a paginated response: {}".format(r.text))
            page_response = r.json()  
            for p_response in page_response:  
                assignments_found_thus_far.append(p_response)

    return assignments_found_thus_far


def submission_for_assignment_by_user(course_id, assignment_id, user_id):
    # return the submission information for a single user's assignment for a specific course as a dict
    #
    # Use the Canvas API to get a user's submission for a course for a specific assignment
    # GET /api/v1/courses/:course_id/assignments/:assignment_id/submissions/:user_id
    url = "{0}/courses/{1}/assignments/{2}/submissions/{3}".format(baseUrl, course_id, assignment_id, user_id)
    if Verbose_Flag:
        print("url: {}".format(url))

    #extra_parameters={'student_ids[]': 'all'}
    #r = requests.get(url, params=extra_parameters, headers = header)
    r = requests.get(url, headers = header)
    if Verbose_Flag:
        print("result of getting submissions: {}".format(r.text))
        
    if r.status_code == requests.codes.ok:
        page_response=r.json()
        if Verbose_Flag:
            print("page_response: " + str(page_response))
        return page_response
    else:
        return dict()

def main():
    global Verbose_Flag

    parser = optparse.OptionParser()

    parser.add_option('-v', '--verbose',
                      dest="verbose",
                      default=False,
                      action="store_true",
                      help="Print lots of output to stdout"
    )

    parser.add_option('-C', '--containers',
                      dest="containers",
                      default=False,
                      action="store_true",
                      help="for the container enviroment in the virtual machine"
    )

    parser.add_option('-t', '--testing',
                      dest="testing",
                      default=False,
                      action="store_true",
                      help="execute test code"
    )

    parser.add_option('-s', '--submit',
                      dest="submit",
                      default=False,
                      action="store_true",
                      help="submit resulting DOCX file as the student"
    )

    parser.add_option("--config", dest="config_filename",
                      help="read configuration from FILE", metavar="FILE")

    options, remainder = parser.parse_args()

    Verbose_Flag=options.verbose
    if Verbose_Flag:
        print("ARGV      : {}".format(sys.argv[1:]))
        print("VERBOSE   : {}".format(options.verbose))
        print("REMAINING : {}".format(remainder))
        print("Configuration file : {}".format(options.config_filename))

    initialize(options)
    if (len(remainder) < 2):
        print("Insuffient arguments - must provide course_id assignment_id")
        sys.exit()
    else:
        course_id=remainder[0]
        assignment_id=int(remainder[1])
       
    if Verbose_Flag:
        print("course_id={0}, assignment_id={1}".format(course_id, assignment_id))

    all_assignments=list_assignments(course_id)
    # check that the assignment is in this course
    assignment=False
    for a in all_assignments:
        if a['id'] == assignment_id:
            assignment=a
            break

    if not assignment:
        print("assignment_id={} not found in this course".format(assignment_id))
        print("The assignments are: {}".format(all_assignments))
        return

    enrollments=users_in_course(course_id)
    # computer a set of user_id for the students in the course
    student_ids=set()
    for e in enrollments:
        if e['type'] == 'StudentEnrollment':
            student_ids.add(e['user_id'])

    print("student_ids={}".format(student_ids))
    for s in student_ids:
        submission_info=submission_for_assignment_by_user(course_id, assignment_id, s)
        # when a submission has been answered with a text ebtry, the text is placed in the body isn HTML format.
        if not submission_info['body']:
            continue
        if submission_info['body']:
            student=get_user(s)
            print("{0}: submission_info={1}".format(student['sortable_name'], submission_info))

            document = Document('blank.docx') # start with a blank A4 page document
            style = document.styles['Normal']
            font = style.font
            font.name = 'Garamond'
            font.size = Pt(12)

            paragraph_format = document.styles['Normal'].paragraph_format
            paragraph_format.space_before = Pt(24)

            about_submission_text="submission by '{0}' (canvas user_id={1}) in course={2} for assignment={3}, submitted_at={4}".format(student['name'], s, course_id, assignment_id, submission_info['submitted_at'])
            p = document.add_paragraph(about_submission_text)
            submission_body=submission_info['body']
            
            p=document.add_paragraph('')
            font = p.add_run('Submission in HTML:').font
            font.highlight_color = WD_COLOR_INDEX.YELLOW
            p = document.add_paragraph(submission_body)

            p=document.add_paragraph('')
            font = p.add_run('Submission:').font
            font.highlight_color = WD_COLOR_INDEX.YELLOW

            paragraphs=get_paragraphs_from_HTML(submission_body)
            for bp in paragraphs:
                p = document.add_paragraph(bp)

            document_name="{0}-{1}-{2}-{3}.docx".format(student['name'], course_id, assignment_id, submission_info['submitted_at'])
            document.save(document_name)
       
            if not options.submit: # if the option to submit is not given, then just generate the DOCX files (locally)
                continue

            status=enable_docx_submission_for_assignment(course_id, assignment_id, assignment)
            if not status:
                print("unable to enable for docx submission")
                continue

            Verbose_Flag=True
            print("about to submit a file")
            if not options.testing:
                submit_file_for_assignment(course_id, assignment_id, s, document_name)
            Verbose_Flag=False

            submission_info2=submission_for_assignment_by_user(course_id, assignment_id, s)
            print("submission_info after submission={}".format(submission_info2))

if __name__ == "__main__": main()

