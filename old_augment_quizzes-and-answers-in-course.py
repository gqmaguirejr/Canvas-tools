#!/usr/bin/python3
# -*- coding: utf-8 -*-
#
# ./augment_quizzes-and-answers-in-course.py course_id
#
# Input:  XLSX spreadsheet with quizzes in course, with name of the form quizzes-<course_id>.xlsx
# Output: augmented XLSX spreadsheet with quizzes in course, with name of the form quizzes-<course_id>-augmented.xlsx
#
#
# with the option "-v" or "--verbose" you get lots of output - showing in detail the operations of the program
#
# Can also be called with an alternative configuration file:
# ./list_your_courses.py --config config-test.json
#
# Example:
# ./augment_quizzes-and-answers-in-course.py 11
#
# ./augment_quizzes-and-answers-in-course.py --config config-test.json 11
#
# 
# documentation about using xlsxwriter to insert images can be found at:
#   John McNamara, "Example: Inserting images into a worksheet", web page, 10 November 2018, https://xlsxwriter.readthedocs.io/example_images.html
#
# G. Q. Maguire Jr.
#
# based on earlier quizzes-and-answers-in-course.py
#
# 2019.01.05
#

import requests, time
import pprint
import optparse
import sys
import json
import os
from pathlib import Path

# Use Python Pandas to create XLSX files
import pandas as pd

# use lxml to access the HTML content
from lxml import html

# use the request pacek to get the HTML give an URL
import requests

# import to be able to compute digests for blank_id
import hashlib

# to be able to read the printed dict() in the spreadsheets
from ast import literal_eval


# use the Docment library to be able to create the .docx file
from docx import Document
#from docx.shared import Inches
from docx.shared import Cm
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx.shared import RGBColor

#############################
###### EDIT THIS STUFF ######
#############################

global baseUrl	# the base URL used for access to Canvas
global header	# the header for all HTML requests
global payload	# place to store additionally payload when needed for options to HTML requests

# Based upon the options to the program, initialize the variables used to access Canvas gia HTML requests
def initialize(options):
    global baseUrl, header, payload

    # styled based upon https://martin-thoma.com/configuration-files-in-python/
    if options.config_filename:
        config_file=options.config_filename
    else:
        config_file='config.json'

    try:
        with open(config_file) as json_data_file:
            configuration = json.load(json_data_file)
            access_token=configuration["canvas"]["access_token"]
            baseUrl="https://"+configuration["canvas"]["host"]+"/api/v1"

            header = {'Authorization' : 'Bearer ' + access_token}
            payload = {}
    except:
        print("Unable to open configuration file named {}".format(config_file))
        print("Please create a suitable configuration file, the default name is config.json")
        sys.exit()
    
def make_dir_for_urls(url, target_dir):
    global Verbose_Flag
    # the URLs have the form: https://canvas.kth.se/courses/11/quizzes/39141/history?quiz_submission_id=759552&version=1
    # remove prefix
    prefix="courses/"
    prefix_offset=url.find(prefix)
    if prefix_offset > 0:
        url_tail=url[prefix_offset+len(prefix):]
        parts_of_path=url_tail.split('/')
        if Verbose_Flag:
            print(parts_of_path)
        course_id=parts_of_path[0]
        quiz_id=parts_of_path[2]
        quiz_submission_part=parts_of_path[3].split('=')
        if Verbose_Flag:
            print(quiz_submission_part)
        quiz_submission_id=quiz_submission_part[1].split('&')[0]
        if Verbose_Flag:
            print(quiz_submission_id)
        dir_to_create="{0}/{1}/{2}/{3}".format(target_dir, course_id, quiz_id, quiz_submission_id)
        print("Creating directory: {}".format(dir_to_create))
        Path(dir_to_create).mkdir(parents=True, exist_ok=True)
        return dir_to_create

def dir_name_for_urls(url, target_dir):
    global Verbose_Flag
    # the URLs have the form: https://canvas.kth.se/courses/11/quizzes/39141/history?quiz_submission_id=759552&version=1
    # remove prefix
    prefix="courses/"
    prefix_offset=url.find(prefix)
    if prefix_offset > 0:
        url_tail=url[prefix_offset+len(prefix):]
        parts_of_path=url_tail.split('/')
        if Verbose_Flag:
            print(parts_of_path)
        course_id=parts_of_path[0]
        quiz_id=parts_of_path[2]
        quiz_submission_part=parts_of_path[3].split('=')
        if Verbose_Flag:
            print(quiz_submission_part)
        quiz_submission_id=quiz_submission_part[1].split('&')[0]
        if Verbose_Flag:
            print(quiz_submission_id)
        dir_to_create="{0}/{1}/{2}/{3}".format(target_dir, course_id, quiz_id, quiz_submission_id)
        return dir_to_create


def compute_canvas_blank_id_digest(bid):
    m=hashlib.md5()
    s1="dropdown,{},instructure-key".format(bid)
    m.update(s1.encode('utf-8'))
    return m.hexdigest()

def fill_common_blank_ids(num):
    global common_blank_ids

    for i in range(0, 10):
        s1="a{}".format(i)
        common_blank_ids[compute_canvas_blank_id_digest(s1)]=s1


# produces '443f6a7201fd1b9c37119a9d1aa28776' - which is indeed the first of the values above
# m=hashlib.md5()
# m.update("dropdown,a1,instructure-key".encode('utf-8'))
# m.hexdigest()
# produces 'd5157c76409909bc0db8eea11b8b37d8' - wicih is the second value above

def cleanup_question_text(text):
    text=text.replace('&nbsp;', ' ')
    text=text.replace('<br>', '\n')
    text=text.replace('<span>', '')
    text=text.replace('</span>', '')
    text=text.replace('<strong>', '')
    text=text.replace('</strong>', '')
    text=text.replace('<p> </p>', ' ')

    return text


# functions for working with the DOCX report document
def docx_create_element(name):
    return OxmlElement(name)

def docx_create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def docx_add_page_number(paragraph):
    # https://python-docx.readthedocs.io/en/latest/api/enum/WdAlignParagraph.html
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    page_run = paragraph.add_run()
    fldChar1 = docx_create_element('w:fldChar')
    docx_create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = docx_create_element('w:instrText')
    docx_create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = docx_create_element('w:fldChar')
    docx_create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_run._r.append(fldChar1)
    page_run._r.append(instrText)
    page_run._r.append(fldChar2)

def main():
    global Verbose_Flag
    global question_type_stats
    global common_blank_ids

    default_picture_size=128

    parser = optparse.OptionParser()

    parser.add_option('-v', '--verbose',
                      dest="verbose",
                      default=False,
                      action="store_true",
                      help="Print lots of output to stdout"
    )
    parser.add_option("--config", dest="config_filename",
                      help="read configuration from FILE", metavar="FILE")
    
    options, remainder = parser.parse_args()

    Verbose_Flag=options.verbose
    if Verbose_Flag:
        print("ARGV      : {}".format(sys.argv[1:]))
        print("VERBOSE   : {}".format(options.verbose))
        print("REMAINING : {}".format(remainder))
        print("Configuration file : {}".format(options.config_filename))

    initialize(options)

    if (len(remainder) < 1):
        print("Insuffient arguments - must provide course_id\n")
    else:
        course_id=remainder[0]
        if not course_id:
            return

    target_dir="./Quiz_Submissions"
    question_id_prefix="question_"

    input_file="quizzes-{}.xlsx".format(course_id)
    excel_Sheet_names = (pd.ExcelFile(input_file)).sheet_names
    if 'Course' in excel_Sheet_names:
        course_df = pd.read_excel(open(input_file, 'rb'), sheet_name='Course')        
        course_code = course_df.iloc[0]['course_code']
        print("course_code={}".format(course_code))

    if 'Quizzes' not in excel_Sheet_names:
        print("Spreadsheeting missing 'Quizzes' sheet, unable to process this file")
        return
    quizzes_df = pd.read_excel(open(input_file, 'rb'), sheet_name='Quizzes')
    quiz_info=dict()
    quiz_list=[]
    for index, row in  quizzes_df.iterrows():
        quiz_info[row['id']]=row['title']

    print("quiz_info={}".format(quiz_info))

    # collect the blank_ids and compute the hashes for them
    #
    common_blank_ids=dict()     #  stores the hashes
    blank_ids=set()		# all the blank_id that appear in any questions
    quizzes_correct_multi_blank_answers=dict()	# index by quiz_id and blank_id
    for quiz in quiz_info:
        sheet_name="{}".format(quiz)
        print("processing quiz={}".format(sheet_name))
        quiz_instance = pd.read_excel(open(input_file, 'rb'), sheet_name=sheet_name)

        for index, row in  quiz_instance.iterrows():        
            # get the question type, if it is fill_in_multiple_blanks_question, then collect the blank names
            q_type=row['question_type']
            if Verbose_Flag:
                print("q_type={}".format(q_type))
            if q_type in ['fill_in_multiple_blanks_question']:
                # [{'weight': 100, 'id': 12882, 'text': 'traffic', 'blank_id': 'a0'}, {'weight': 100, 'id': 63085, 'text': 'energy consumption', 'blank_id': 'a0'}, {'weight': 100, 'id': 51091, 'text': 'energy consumption', 'blank_id': 'a1'}, {'weight': 100, 'id': 39757, 'text': 'traffic', 'blank_id': 'a1'}]
                correct_multi_blank_answers=dict()	# index by blank_id

                q_answers=row['answers']
                if Verbose_Flag:
                    print("q_answers={0}, type(q_answers)={1}".format(q_answers, type(q_answers)))
                q_answers=literal_eval(row['answers'])
                for ans in q_answers:
                    b_id=ans['blank_id']
                    if Verbose_Flag:
                        print("b_id={}".format(b_id))
                    blank_ids.add(b_id)
                    if ans['weight'] == 100:
                        if correct_multi_blank_answers.get(b_id, None):
                            correct_multi_blank_answers[b_id].append(ans['text'])
                        else:
                            correct_multi_blank_answers[b_id]=[ans['text']]

                q_id=row['id']
                #print("q_id={0}, type(q_id)={1}".format(q_id, type(q_id)))
                quizzes_correct_multi_blank_answers[q_id]=correct_multi_blank_answers

    if Verbose_Flag:
        print("blank_ids={}".format(sorted(blank_ids)))
    # if any blank_id is missing from those with entries in common_blank_ids add the mapping
    for b_id in blank_ids:
        common_blank_ids[compute_canvas_blank_id_digest(b_id)]=b_id

    if Verbose_Flag:
        print("common_blank_ids={}".format(common_blank_ids))

    print("quizzes_correct_multi_blank_answers={}".format(quizzes_correct_multi_blank_answers))

    incorrect_answers=dict()    # values a set under the question_id
    incorrect_answers_multiple_blanks=dict() #  values are stored as quiest_id: {blank_id_hash: {value set}}

    quiz_submissions_all=dict()
    for quiz in quiz_info:
        sheet_name="s_{}".format(quiz)
        quiz_submissions_all[quiz] = pd.read_excel(open(input_file, 'rb'), sheet_name=sheet_name)

        # add an empty column to remeber the questions that were ask
        quiz_submissions_all[quiz]['questions_asked']=''
        quiz_submissions_all[quiz]['questions_types']=''
        quiz_submissions_all[quiz]['answer_correctness']=''

        for index, row in  quiz_submissions_all[quiz].iterrows():
            # make a directory for the submissions from result_url
            # the URLs have the form: https://canvas.kth.se/courses/11/quizzes/39141/history?quiz_submission_id=759552&version=1
            url=row['result_url']
            questions_asked=dict()
            questions_types=dict()
            answer_correctness=dict()
            version_string=''
            dir_name=dir_name_for_urls(url, target_dir)
            for file in os.listdir(dir_name):
                version_string=''
                if file.endswith(".html"):
                    file_name=os.path.join(dir_name, file)
                    html_string=''
                    try:
                        with open(file_name) as html_file:
                            html_string = html_file.read()
                    except:
                        print("Unable to open file named {}".format(file_name))

                    print("\nProcessing file: {}".format(file_name))
                    version_prefix='_v'
                    version_offset=file_name.rindex(version_prefix)+2
                    if version_offset:
                        version_string=file_name[version_offset:-5]
                        if Verbose_Flag:
                            print("version_string={}".format(version_string))

                    if len(html_string) == 0:
                        print("No contents for file={}".format(file_name))
                        continue
                    document = html.document_fromstring(html_string)
                    display_questions = document.xpath('//*[contains(@class, "display_question")]')
                    # Can have a correct answer:
                    # question id=question_274025 class=display_question question multiple_answers_question  correct   
                    # question id=question_274030 class=display_question question true_false_question  correct   
                    # question id=question_274031 class=display_question question multiple_answers_question  correct   
                    # question id=question_274032 class=display_question question short_answer_question  correct   

                    # Can have partially correct answer:
                    # question id=question_274033 class=display_question question multiple_answers_question     partial_credit
                    #
                    # Can have an incorrect answer:
                    # question id=question_274222 class=display_question question short_answer_question   incorrect  
                    # 
                    # Remember the questions ask in this attempt
                    question_asked_this_attempt=list()
                    questions_types_this_attempt=list()
                    answer_correctness_this_attempt=list()

                    for a_question in display_questions:
                        question_type=None
                        a_question_id=a_question.attrib['id']
                        a_question_class=a_question.attrib['class']
                        if Verbose_Flag:
                            print("question id={0} class={1}".format(a_question_id, a_question_class))
                        question_id=int(a_question_id[len(question_id_prefix):])
                        question_asked_this_attempt.append(question_id)
                        if 'partial_credit' in a_question_class:
                            answer_correctness_this_attempt.append('partial_credit')
                        elif 'incorrect' in a_question_class:
                            answer_correctness_this_attempt.append('incorrect')
                        else:
                            # otherwise assume it is correct
                            answer_correctness_this_attempt.append('correct')

                        if 'multiple_answers_question' in a_question_class:
                            question_type='multiple_answers_question'
                        elif 'short_answer' in  a_question_class:
                            question_type='short_answer'
                        elif 'true_false_question' in  a_question_class:
                            question_type='true_false_question'
                        elif 'fill_in_multiple_blanks_question' in  a_question_class:
                            question_type='fill_in_multiple_blanks_question'
                        elif 'multiple_choice_question' in  a_question_class:
                            question_type='multiple_choice_question'
                        elif 'matching_question' in  a_question_class:
                            question_type='matching_question'
                        else:
                            question_type=a_question_class

                        questions_types_this_attempt.append(question_type)

                        #  note that this is a relative query, hence the leading "." - so we only look at answers to this question
                        inputs = a_question.xpath('.//*[@class="question_input"]')
                        #<input type="text" name="question_274239" value="effective" class="question_input" autocomplete="off" readonly="&quot;readonly&quot;" aria-label="Fill in the blank answer" aria-describedby="user_answer_NaN_arrow" style="width: 85.5px;">
                        # version_string
                        for a_number, input in enumerate(inputs):
                            if Verbose_Flag:
                                print("a_number={}".format(a_number))
                            input_id=input.attrib.get('id', None)
                            input_type=input.attrib.get('type', None)
                            input_name=input.attrib.get('name', None)
                            input_value=input.attrib.get('value', None)
                            input_describedby=input.attrib.get('aria-describedby', None)
                            if not input_value: # skip case of no answer
                                continue

                            if 'incorrect' in a_question_class and input_type=='text':
                                print("input_id={0}, input_type={1},input_name={2},input_value={3},input_describedby={4}, question_type={5}".format(input_id, input_type, input_name, input_value, input_describedby, question_type))
                                split_input_name=input_name.split('_')
                                question_id=int(split_input_name[1])
                                if len(split_input_name) > 2:
                                    print("extention to input_name for question_id={0} of {1}".format(question_id, split_input_name))
                                if question_type == 'fill_in_multiple_blanks_question':
                                    # note that quseions with multiple blanks have gensymn'd extensions to the name:
                                    # input_name=question_274226_443f6a7201fd1b9c37119a9d1aa28776 - has the blank_id='a0'
                                    # input_name=question_274226_d5157c76409909bc0db8eea11b8b37d8 - has the blank_id='a1'
                                    # However, I do not know (yet) how to know that the above are the mappings (except by examinging the question manually)
                                    # The hash is computed as specfied in Canvas app/models/assessment_question.rb
                                    # def self.variable_id(variable)
                                    #   Digest::MD5.hexdigest(["dropdown", variable, "instructure-key"].join(","))
                                    # end
                                    # so we see it putting the variable between two string separated by commas before computing the MD5 hash
                                    # import hashlib
                                    # m=hashlib.md5()
                                    # m.update("dropdown,a0,instructure-key".encode('utf-8'))
                                    # m.hexdigest()
                                    # produces '443f6a7201fd1b9c37119a9d1aa28776' - which is indeed the first of the values above
                                    # m=hashlib.md5()
                                    # m.update("dropdown,a1,instructure-key".encode('utf-8'))
                                    # m.hexdigest()
                                    # produces 'd5157c76409909bc0db8eea11b8b37d8' - wicih is the second value above
                                    # NB you have to create a new shash object (m) otherwise the update data is just appended to what was already there

                                    current_incorrect_answers=incorrect_answers_multiple_blanks.get(question_id, None)
                                    blank_id_hash=split_input_name[2]
                                    # If we know the hashid value, then use the actual blank_id
                                    blank_id=common_blank_ids.get(blank_id_hash, None)
                                    if blank_id:
                                        blank_id_hash=blank_id
                                    if current_incorrect_answers:
                                        c_answers=quizzes_correct_multi_blank_answers.get(question_id, None)
                                        b_c_answers=c_answers.get(blank_id_hash, None)
                                        incorrect_answers_for_blank=current_incorrect_answers.get(blank_id_hash, None)
                                        if incorrect_answers_for_blank:
                                            if input_value not in b_c_answers:
                                                incorrect_answers_for_blank.add(input_value) #  note that the add occurs in place
                                        else:
                                            if input_value not in b_c_answers:
                                                incorrect_answers_multiple_blanks[question_id][blank_id_hash]= {input_value}
                                    else:
                                        incorrect_answers_multiple_blanks[question_id]={blank_id_hash: {input_value}}
                                else: #  not a multiple blank question
                                    # Note that we need to filter out the correct answers forthe parts of the questions
                                    # since for a partially correct answer - some of the answers could be correct, and the get added to the list currently!
                                    # This needs to be fixed.
                                    current_incorrect_answers=incorrect_answers.get(question_id, None)
                                    if current_incorrect_answers:
                                        current_incorrect_answers.add(input_value) #  note that the add occurs in place
                                        print("first case: input_value={0}, incorrect_answers[{1}]={2}".format(input_value, question_id, incorrect_answers[question_id]))
                                    else:
                                        incorrect_answers[question_id]={input_value}
                                        print("second case: input_value={0}, incorrect_answers[{1}]={2}".format(input_value, question_id, incorrect_answers[question_id]))



                    questions_asked[version_string]=question_asked_this_attempt
                    questions_types[version_string]=questions_types_this_attempt
                    answer_correctness[version_string]=answer_correctness_this_attempt


            if questions_asked:
                print("questions_asked={}".format(questions_asked))
                quiz_submissions_all[quiz].at[index, 'questions_asked']=questions_asked

            if questions_types:
                print("questions_types={}".format(questions_types))
                quiz_submissions_all[quiz].at[index, 'questions_types']=questions_types

            if answer_correctness:
                print("answer_correctness={}".format(answer_correctness))
                quiz_submissions_all[quiz].at[index, 'answer_correctness']=answer_correctness


    # the following was inspired by the section "Using XlsxWriter with Pandas" on http://xlsxwriter.readthedocs.io/working_with_pandas.html
    # set up the output write
    writer = pd.ExcelWriter('quizzes-'+course_id+'-augmented.xlsx', engine='xlsxwriter')
    quizzes_df.to_excel(writer, sheet_name='Quizzes')

    quiz_report_document = Document()
    quiz_report_section = quiz_report_document.sections[0]

    # Set for A4 page size
    quiz_report_section.page_height = Mm(297)
    quiz_report_section.page_width = Mm(210)

    quiz_report_header = quiz_report_section.header
    quiz_report_paragraph = quiz_report_header.paragraphs[0]
    quiz_report_paragraph.text = "{0}\t{1}\tQuiz report".format(course_code, course_id)
    quiz_report_paragraph.style = quiz_report_document.styles["Header"]

    # Add a footer with the page number
    docx_add_page_number(quiz_report_document.sections[0].footer.paragraphs[0])

    question_types_of_interest=['short_answer_question', 'fill_in_multiple_blanks_question']
    # Add top level heading (title of the report)
    quiz_report_document.add_heading("Quiz questions and answers: Focus on questions of types: {0}".format(question_types_of_interest), 0)

    author_string="augment_quizzes-and-answers-in-course.py {}".format(course_id)
    p = quiz_report_document.add_paragraph('Output of ')
    current_font=p.add_run(author_string).font
    current_font.color.rgb = RGBColor(0xFF, 0x00, 0xFF)
    current_font.name = 'Consolas'

    quiz_questions_to_quiz=dict() # for each question_id record what quiz it appears in

    for qi, quiz in enumerate(quiz_info):
        # copy over the quiz data
        sheet_name="{}".format(quiz)
        print("copying quiz={}".format(sheet_name))
        quiz_instance = pd.read_excel(open(input_file, 'rb'), sheet_name=sheet_name)

        if qi > 0:
            quiz_report_document.add_page_break()
        quiz_report_document.add_heading("{0}: {1}".format(quiz, quiz_info[quiz]), 1)


        # create a column to put the incorrce answers into
        quiz_instance['incorrect answers']=''

        for index, row in  quiz_instance.iterrows():        
            if quiz_questions_to_quiz.get(row['id'], None):
                print("question_id={0} already has a quiz_id {1} associated with it".format(row['id'], row['quiz_id']))
            else:
                quiz_questions_to_quiz[row['id']]=row['quiz_id']

            #question_type: short_answer_question
            #question_id: 379340 (the id of a question within a given quiz)
            #question_name: Philosophic basis of anthropocentric thinking
            #question_text: Putting people and their well-being first is the philosophic basis of _______&nbsp;thinking.
            # only add limited number of types of questions to the report
            if row['question_type'] in question_types_of_interest:
                add_to_report=True
            else:
                add_to_report=False

            if add_to_report:
                #rt="question_name: {}".format(row['question_name'])
                rt="{}".format(row['question_name'])
                quiz_report_document.add_heading(rt, level=2)

                rt="question_type: {}".format(row['question_type'])
                p = quiz_report_document.add_paragraph(rt)

                rt="question_id: {}".format(row['id'])
                p = quiz_report_document.add_paragraph(rt)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                rt="question_text: {}".format(cleanup_question_text(row['question_text']))
                p = quiz_report_document.add_paragraph(rt)

                rt="answers: "
                p = quiz_report_document.add_paragraph(rt)

                if row['question_type'] in ['fill_in_multiple_blanks_question']:
                    rt_ia=dict()
                    rt_a=literal_eval(row['answers'])
                    for ia in rt_a:
                        #{'weight': 100, 'id': 12882, 'text': 'traffic', 'blank_id': 'a0'}, 
                        rt_i=ia['blank_id']
                        if rt_ia.get(rt_i, None):
                            rt_ia[rt_i].append(ia['text'])
                        else:
                            rt_ia[rt_i]=[ia['text']]
                    #print("rt_ia={}".format(rt_ia))
                    rt="{}".format(rt_ia)
                    p.add_run(rt)
                elif row['question_type'] in ['short_answer_question']:
                    rt_ia=list()
                    rt_a=literal_eval(row['answers'])
                    for ia in rt_a:
                        if ia['weight'] == 100:
                            rt_ia.append(ia['text'])
                    #rt="{}".format(row['answers'])
                    rt="{}".format(rt_ia)
                    p.add_run(rt)
                else:
                    print("other case={}".format(row['answers']))
                    rt="{}".format(row['answers'])
                    p.add_run(rt)

            if row['id'] in incorrect_answers:
                quiz_instance.at[index, 'incorrect answers']=incorrect_answers[row['id']]
                if add_to_report:
                    rt="incorrect answers: {}".format(incorrect_answers[row['id']])
                    p = quiz_report_document.add_paragraph(rt)

            if row['id'] in incorrect_answers_multiple_blanks:
                quiz_instance.at[index, 'incorrect answers']=incorrect_answers_multiple_blanks[row['id']]
                if add_to_report:
                    rt="incorrect answers: {}".format(incorrect_answers_multiple_blanks[row['id']])
                    p = quiz_report_document.add_paragraph(rt)


        # save updated quiz_instance
        quiz_instance.to_excel(writer, sheet_name=sheet_name)

        # save the aumented quiz submission information
        sheet_name="s_{}".format(quiz)
        print("saving quiz submissions={}".format(sheet_name))
        quiz_submissions_all[quiz].to_excel(writer, sheet_name=sheet_name)

    # Close the Pandas Excel writer and output the Excel file.
    writer.save()

    for q in incorrect_answers:
        try:
            print("Incorrect answers for {0} on quiz {1}: {2}".format(q, quiz_questions_to_quiz[q], incorrect_answers[q]))
        except:
            print("Invalid question id {0} (not in the list of questions - but has incorrect answers={1}".format(q, incorrect_answers[q]))

    for q in incorrect_answers_multiple_blanks:
        try:
            print("Incorrect answers for {0} on quiz {1}: {2}".format(q, quiz_questions_to_quiz[q], incorrect_answers_multiple_blanks[q]))
        except:
            print("Invalid question id {0} (not in the list of questions - but has incorrect answers={1}".format(q, incorrect_answers_multiple_blanks[q]))

    quiz_report_document.save('quizzes-'+course_id+'-augmented.docx')
    
if __name__ == "__main__": main()

